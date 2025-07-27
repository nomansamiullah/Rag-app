import os
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import mimetypes
from datetime import datetime

# Document processors
import PyPDF2
from docx import Document
import pandas as pd
from openpyxl import load_workbook

# Text splitting
from langchain.text_splitter import RecursiveCharacterTextSplitter

from utils.logger import setup_logger
from config import Config

logger = setup_logger(__name__)


from typing import List, Dict, Any
from utils.logger import setup_logger

logger = setup_logger(__name__)


class FileProcessor:
    def __init__(self):
        """Initialize file processor with text splitter"""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        
        # Supported file types
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # Will try docx processor
            '.txt': self._process_txt,
            '.md': self._process_txt,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
        }
        
        logger.info("File processor initialized")
    
    def get_file_info(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Extract file metadata"""
        try:
            file_stats = os.stat(file_path)
            file_extension = Path(original_filename).suffix.lower()
            
            # Calculate file hash for duplicate detection
            file_hash = self._calculate_file_hash(file_path)
            
            return {
                'filename': original_filename,
                'file_extension': file_extension,
                'file_size': file_stats.st_size,
                'file_hash': file_hash,
                'upload_timestamp': datetime.now().isoformat(),
                'mime_type': mimetypes.guess_type(original_filename)[0] or 'unknown',
                'is_supported': file_extension in self.supported_types
            }
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            raise
    
    def process_file(self, file_path: str, original_filename: str, 
                    custom_metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """Process file and return chunks with metadata"""
        try:
            # Get file info
            file_info = self.get_file_info(file_path, original_filename)
            
            if not file_info['is_supported']:
                raise ValueError(f"Unsupported file type: {file_info['file_extension']}")
            
            # Extract text content
            file_extension = file_info['file_extension']
            processor_func = self.supported_types[file_extension]
            
            logger.info(f"Processing {original_filename} with {processor_func.__name__}")
            text_content = processor_func(file_path)
            
            if not text_content or not text_content.strip():
                raise ValueError("No text content extracted from file")
            
            # Split into chunks
            chunks = self._create_chunks(text_content)
            
            # Create chunk documents with metadata
            chunk_documents = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    **file_info,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'chunk_size': len(chunk),
                    'source_type': 'file_upload'
                }
                
                # Add custom metadata if provided
                if custom_metadata:
                    chunk_metadata.update(custom_metadata)
                
                chunk_documents.append({
                    'text': chunk,
                    'metadata': chunk_metadata
                })
            
            logger.info(f"Processed {original_filename}: {len(chunks)} chunks created")
            return chunk_documents
            
        except Exception as e:
            logger.error(f"Error processing file {original_filename}: {str(e)}")
            raise
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash of file for duplicate detection"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def _create_chunks(self, text: str) -> List[str]:
        """Split text into chunks"""
        chunks = self.text_splitter.split_text(text)
        # Filter out very short chunks
        chunks = [chunk for chunk in chunks if len(chunk.strip()) > Config.MIN_CHUNK_SIZE]
        return chunks
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise
    
    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise
    
    def _process_txt(self, file_path: str) -> str:
        """Extract text from TXT/MD files"""
        try:
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode file with any supported encoding")
            
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise
    
    def _process_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert dataframe to readable text
            text_content = f"CSV Data Summary:\n"
            text_content += f"Columns: {', '.join(df.columns.tolist())}\n"
            text_content += f"Total Rows: {len(df)}\n\n"
            
            # Add column descriptions
            text_content += "Column Information:\n"
            for col in df.columns:
                col_info = f"- {col}: {df[col].dtype}"
                if df[col].dtype in ['object', 'string']:
                    unique_count = df[col].nunique()
                    text_content += f" ({unique_count} unique values)"
                text_content += "\n"
            
            # Add sample data
            text_content += f"\nSample Data (first 10 rows):\n"
            text_content += df.head(10).to_string(index=False)
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_content += f"\n\nNumeric Summary:\n"
                text_content += df[numeric_cols].describe().to_string()
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error processing CSV: {str(e)}")
            raise
    
    def _process_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            # Load workbook to get sheet names
            workbook = load_workbook(file_path, read_only=True)
            sheet_names = workbook.sheetnames
            
            text_content = f"Excel File Content:\n"
            text_content += f"Sheets: {', '.join(sheet_names)}\n\n"
            
            # Process each sheet
            for sheet_name in sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    text_content += f"--- Sheet: {sheet_name} ---\n"
                    text_content += f"Columns: {', '.join(df.columns.tolist())}\n"
                    text_content += f"Rows: {len(df)}\n\n"
                    
                    # Add sample data
                    if len(df) > 0:
                        text_content += f"Sample Data:\n"
                        text_content += df.head(5).to_string(index=False)
                        text_content += "\n\n"
                    
                except Exception as e:
                    logger.warning(f"Error processing sheet {sheet_name}: {str(e)}")
                    continue
            
            workbook.close()
            return text_content
            
        except Exception as e:
            logger.error(f"Error processing Excel: {str(e)}")
            raise




# use this without file uplaod
#class FileProcessor:
#     def __init__(self):
#         """Simple file processor without langchain"""
#         logger.info("Simple file processor initialized")
    
#     def split_text(self, text: str, chunk_size: int = 1000) -> List[str]:
#         """Simple text splitting without langchain"""
#         chunks = []
#         for i in range(0, len(text), chunk_size):
#             chunk = text[i:i+chunk_size]
#             if len(chunk.strip()) > 100:  # Only keep meaningful chunks
#                 chunks.append(chunk.strip())
#         return chunks